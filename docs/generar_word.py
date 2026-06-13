#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el documento Word de Codex FCT26."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colores ─────────────────────────────────────────────────────────────────
DARK_BG    = RGBColor(0x0F, 0x17, 0x2A)
DARK_BG2   = RGBColor(0x1E, 0x29, 0x3B)
PRIMARY    = RGBColor(0x38, 0xBD, 0xF8)
ACCENT     = RGBColor(0x81, 0x8C, 0xF8)
GREEN      = RGBColor(0x4A, 0xDE, 0x80)
RED        = RGBColor(0xF8, 0x71, 0x71)
WHITE      = RGBColor(0xF8, 0xFA, 0xFC)
MUTED      = RGBColor(0x94, 0xA3, 0xB8)
CODE_BG    = RGBColor(0x0D, 0x14, 0x21)
CODE_TEXT  = RGBColor(0xCB, 0xD5, 0xE1)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)


def add_page_border(doc):
    """Add a subtle left border accent to every section."""
    pass  # We'll handle via styles


def set_doc_margins(doc, top=2, bottom=2, left=2.5, right=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)


def set_run_font(run, size_pt, bold=False, color=None, italic=False, font_name='Inter'):
    run.font.name    = font_name
    run.font.size    = Pt(size_pt)
    run.font.bold    = bold
    run.font.italic  = italic
    if color:
        run.font.color.rgb = color


def heading(doc, text, level=1, color=PRIMARY):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.name = 'Inter'
    return p


def body(doc, text, color=WHITE, size=10.5, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, color=color, italic=italic)
    p.paragraph_format.space_after = Pt(4)
    return p


def bullet(doc, text, color=WHITE, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_run_font(run, size, color=color)
    p.paragraph_format.space_after = Pt(3)
    return p


def code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(code_text)
    run.font.name  = 'JetBrains Mono'
    run.font.size  = Pt(8.5)
    run.font.color.rgb = CODE_TEXT
    # shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), '0D1421')
    shd.set(qn('w:val'),  'clear')
    pPr.append(shd)
    return p


def divider(doc):
    p = doc.add_paragraph('─' * 60)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.size  = Pt(6)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0x60)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    return p


def section_title_box(doc, text, color_hex='0EA5E9'):
    """Caja de título de sección."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, color_hex)
    cell.width = Cm(16)
    p = cell.paragraphs[0]
    run = p.add_run(f'  {text}')
    run.font.name  = 'Inter'
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()


def info_table(doc, data: list, headers: list, col_widths=None):
    """Tabla de datos estilizada."""
    n_cols = len(headers)
    table  = doc.add_table(rows=1 + len(data), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '0E4F7A')
        p    = cell.paragraphs[0]
        run  = p.add_run(h)
        run.font.name  = 'Inter'
        run.font.size  = Pt(9)
        run.font.bold  = True
        run.font.color.rgb = PRIMARY

    # Data rows
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg  = '0F172A' if r_idx % 2 == 0 else '131F30'
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.name  = 'JetBrains Mono' if c_idx == 1 else 'Inter'
            run.font.size  = Pt(8.5)
            run.font.color.rgb = WHITE if c_idx != 3 else MUTED

    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# DOCUMENTO PRINCIPAL
# ════════════════════════════════════════════════════════════════════════════

doc = Document()
set_doc_margins(doc, top=2.2, bottom=2.2, left=2.8, right=2.5)

# Fondo oscuro del documento
for section in doc.sections:
    sectPr = section._sectPr
    bgCol  = OxmlElement('w:background')
    bgCol.set(qn('w:color'), '0F172A')
    doc._element.insert(0, bgCol)

# ── PORTADA ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run('CODEX')
run.font.name  = 'Inter'
run.font.size  = Pt(42)
run.font.bold  = True
run.font.color.rgb = PRIMARY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = doc.add_paragraph()
run2 = p2.add_run('Plataforma de Aprendizaje Web Interactivo')
run2.font.name  = 'Inter'
run2.font.size  = Pt(16)
run2.font.color.rgb = MUTED
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p3 = doc.add_paragraph()
run3 = p3.add_run('Proyecto FCT26 · Trabajo de Fin de Grado · 2026')
run3.font.name  = 'Inter'
run3.font.size  = Pt(10)
run3.font.color.rgb = MUTED
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p4 = doc.add_paragraph()
run4 = p4.add_run('Angular 21  ·  NestJS  ·  MongoDB Atlas  ·  TypeScript  ·  JWT')
run4.font.name  = 'Inter'
run4.font.size  = Pt(10)
run4.font.color.rgb = ACCENT
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(4):
    doc.add_paragraph()

# Autor
p5 = doc.add_paragraph()
run5 = p5.add_run('Itziar Carnicerarroyo')
run5.font.name  = 'Inter'
run5.font.size  = Pt(12)
run5.font.bold  = True
run5.font.color.rgb = WHITE
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER

p6 = doc.add_paragraph()
run6 = p6.add_run('itziarcarnicerarroyo@gmail.com')
run6.font.name  = 'Inter'
run6.font.size  = Pt(10)
run6.font.color.rgb = MUTED
p6.alignment = WD_ALIGN_PARAGRAPH.CENTER

p7 = doc.add_paragraph()
run7 = p7.add_run('github.com/JRGamer234/Codex-FCT26')
run7.font.name  = 'Inter'
run7.font.size  = Pt(10)
run7.font.color.rgb = PRIMARY
p7.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ── 1. INTRODUCCIÓN ─────────────────────────────────────────────────────────
section_title_box(doc, '1. Introducción', '0E4F7A')

body(doc, '¿Qué es Codex?', color=PRIMARY, size=12, bold=True)
body(doc,
    'Codex es una plataforma educativa web diseñada para enseñar desarrollo frontend de forma '
    'estructurada e interactiva. Los alumnos aprenden HTML, CSS y JavaScript a través de lecciones '
    'que combinan teoría explicada con claridad, ejemplos de código con resaltado sintáctico y quizzes '
    'de evaluación al final de cada tema.',
    size=10.5)

doc.add_paragraph()
body(doc, 'Objetivos del proyecto', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Crear un entorno de aprendizaje web completo y accesible para estudiantes de desarrollo frontend.')
bullet(doc, 'Proporcionar al profesor un panel de control para gestionar alumnos, lecciones y revisar el progreso.')
bullet(doc, 'Implementar una arquitectura moderna full-stack con buenas prácticas de seguridad y escalabilidad.')
bullet(doc, 'Aplicar los conocimientos adquiridos durante el ciclo formativo en un proyecto real y funcional.')

doc.add_paragraph()
body(doc, 'Contexto', color=MUTED, size=9.5, italic=True)
body(doc,
    'El proyecto nace como Trabajo de Fin de Grado (FCT26). La plataforma está completamente '
    'operativa con conexión a base de datos en la nube (MongoDB Atlas), autenticación JWT segura '
    'y separación de roles entre alumno y profesor.',
    size=10.5)

divider(doc)

# ── 2. TECNOLOGÍAS ──────────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '2. Stack tecnológico', '1A3A5C')

body(doc, '2.1 Frontend', color=PRIMARY, size=12, bold=True)
body(doc, 'La interfaz está construida con Angular 21, el framework frontend de Google basado en TypeScript. '
         'Se utilizan componentes standalone (sin NgModules), la nueva API de signals para el estado reactivo '
         'y el servidor de desarrollo Vite para compilación rápida.', size=10.5)

doc.add_paragraph()

info_table(doc,
    data=[
        ['Angular',       '21',         'Framework principal SPA',       'Frontend'],
        ['TypeScript',    '5.x',        'Tipado estático sobre JavaScript','Frontend'],
        ['SCSS',          'Módulos',    'Estilos con variables y nesting', 'Frontend'],
        ['Vite',          'Dev server', 'Compilación rápida en desarrollo','Frontend'],
        ['Prism.js',      '1.x',        'Resaltado de código en lecciones','Frontend'],
        ['RxJS',          '7.x',        'Programación reactiva / observables','Frontend'],
    ],
    headers=['Tecnología', 'Versión', 'Uso', 'Capa']
)

doc.add_paragraph()
body(doc, '2.2 Backend', color=PRIMARY, size=12, bold=True)
body(doc, 'El servidor está construido con NestJS, un framework Node.js de arquitectura modular inspirado en Angular. '
         'Utiliza decoradores TypeScript, inyección de dependencias y una organización clara por módulos.', size=10.5)
doc.add_paragraph()

info_table(doc,
    data=[
        ['NestJS',        '10.x',  'Framework backend principal',             'Backend'],
        ['Node.js',       '24.x',  'Runtime JavaScript del servidor',         'Backend'],
        ['Mongoose',      '8.x',   'ODM para MongoDB',                        'Backend'],
        ['Passport',      '0.7',   'Middleware de autenticación',              'Backend'],
        ['bcrypt',        '5.x',   'Hashing de contraseñas',                  'Backend'],
        ['class-validator','0.14', 'Validación de DTOs con decoradores',       'Backend'],
    ],
    headers=['Tecnología', 'Versión', 'Uso', 'Capa']
)

doc.add_paragraph()
body(doc, '2.3 Base de datos y herramientas', color=PRIMARY, size=12, bold=True)
info_table(doc,
    data=[
        ['MongoDB Atlas',    'M0 Free',    'Base de datos NoSQL en la nube', 'Datos'],
        ['GitHub',           'Git',        'Control de versiones',           'DevOps'],
        ['npm workspaces',   'Monorepo',   'Gestión de paquetes unificada',  'DevOps'],
        ['Angular Proxy',    'Dev',        'Eliminación de CORS en desarrollo','Config'],
        ['dotenv (.env)',     'Variables',  'Configuración de entorno segura','Config'],
    ],
    headers=['Herramienta', 'Tipo', 'Propósito', 'Área']
)

divider(doc)

# ── 3. ARQUITECTURA ─────────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '3. Arquitectura del sistema', '1A3A5C')

body(doc, '3.1 Estructura del repositorio', color=PRIMARY, size=12, bold=True)
doc.add_paragraph()
code_block(doc,
'''Codex-FCT26/
├── frontend/               # Aplicación Angular 21
│   ├── src/app/
│   │   ├── core/           # Servicios, interceptores, guards
│   │   │   ├── services/   # auth, user, lesson, progress, rating
│   │   │   ├── interceptors/ # auth.interceptor (JWT + 401)
│   │   │   └── guards/     # auth.guard, role.guard
│   │   └── features/       # Componentes de página
│   │       ├── dashboard/  login, register, dashboard alumno
│   │       ├── catalog/    catálogo de lecciones
│   │       ├── lesson-detail / lesson-quiz
│   │       ├── profile/    perfil de usuario
│   │       ├── achievements/ sistema de logros
│   │       └── profesor-*/  panel del profesor
│   ├── proxy.conf.json     # Proxy /api/* → localhost:3000
│   └── angular.json
│
├── backend/                # API NestJS
│   ├── src/
│   │   ├── auth/           # Login, register, JWT strategy
│   │   ├── users/          # Perfil, alumnos, crear usuario
│   │   ├── lessons/        # CRUD lecciones
│   │   ├── progress/       # Completar lección, historial
│   │   ├── ratings/        # Valoraciones
│   │   ├── seed.ts         # Poblar base de datos
│   │   └── main.ts         # Configuración del servidor
│   └── .env                # MONGO_URL, JWT_SECRET, PORT
│
├── docs/                   # Documentación del proyecto
└── package.json            # Scripts raíz (npm run dev)''')

doc.add_paragraph()
body(doc, '3.2 Flujo de comunicación', color=PRIMARY, size=12, bold=True)
body(doc,
    'El frontend Angular hace peticiones HTTP relativas (/api/...) que el proxy del servidor de '
    'desarrollo redirige al backend NestJS en localhost:3000. El backend procesa la petición, '
    'consulta MongoDB Atlas y devuelve JSON. Todas las rutas de la API requieren un JWT válido '
    'en la cabecera Authorization (excepto login y register).', size=10.5)

doc.add_paragraph()
code_block(doc,
'Navegador → /api/users/alumnos\n'
'  → Proxy Angular (localhost:4200)\n'
'  → NestJS API (localhost:3000/api/users/alumnos)\n'
'  → JwtAuthGuard verifica token\n'
'  → UsersController → UsersService\n'
'  → MongoDB Atlas (query usuarios + progreso)\n'
'  → JSON response')

divider(doc)

# ── 4. MODELOS DE DATOS ──────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '4. Modelos de datos (MongoDB)', '1A3A5C')

body(doc, '4.1 User', color=PRIMARY, size=12, bold=True)
code_block(doc,
'''@Schema({ timestamps: true })
class User {
  name:     string;          // Nombre completo
  email:    string;          // Único en la colección
  password: string;          // Hash bcrypt (salt 10)
  rol:      'alumno' | 'profesor';  // Rol del usuario
}''')

body(doc, '4.2 Lesson', color=PRIMARY, size=12, bold=True)
code_block(doc,
'''@Schema({ timestamps: true })
class Lesson {
  title:       string;
  description: string;
  level:       'Inicial' | 'Medio' | 'Avanzado';
  category:    'HTML' | 'CSS' | 'JavaScript';
  sections:    Section[];    // { title, content, code }
  questions:   Question[];   // { question, options[], correct }
  createdBy:   string;
}''')

body(doc, '4.3 Progress', color=PRIMARY, size=12, bold=True)
code_block(doc,
'''@Schema({ timestamps: true })
class Progress {
  userId:        ObjectId → User;   // Referencia al alumno
  userName:      string;            // Desnormalizado para consultas rápidas
  lessonId:      string;
  lessonTitle:   string;
  score:         number;            // Aciertos en el quiz
  total:         number;            // Total de preguntas
  completedAt:   Date;
  // Índice único compuesto: (userId, lessonId)
}''')

body(doc, '4.4 Rating', color=PRIMARY, size=12, bold=True)
code_block(doc,
'''@Schema({ timestamps: true })
class Rating {
  userId:      ObjectId → User;
  userName:    string;
  lessonId:    string;
  lessonTitle: string;
  stars:       number;   // 1 - 5
  comment:     string;   // Opcional
}''')

divider(doc)

# ── 5. API ENDPOINTS ─────────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '5. API REST — Endpoints', '1A3A5C')

body(doc,
    'Todos los endpoints tienen el prefijo global /api. Los protegidos requieren la cabecera '
    'Authorization: Bearer <token>. La autorización por rol se verifica en el servidor.', size=10.5)
doc.add_paragraph()

info_table(doc,
    data=[
        ['POST',   '/api/auth/register',      'Registro de nuevo usuario',                      'Público'],
        ['POST',   '/api/auth/login',          'Login · devuelve JWT (7 días)',                  'Público'],
        ['GET',    '/api/users/me',            'Datos del usuario autenticado',                  'Autenticado'],
        ['PATCH',  '/api/users/me',            'Actualizar nombre del perfil',                   'Autenticado'],
        ['GET',    '/api/users/alumnos',       'Lista de alumnos con progreso calculado',        'Profesor'],
        ['POST',   '/api/users/alumnos',       'Crear nuevo alumno (nombre, email, contraseña)', 'Profesor'],
        ['GET',    '/api/lessons',             'Todas las lecciones (título, nivel, categoría)', 'Autenticado'],
        ['GET',    '/api/lessons/:id',         'Lección completa con secciones y quiz',          'Autenticado'],
        ['POST',   '/api/lessons',             'Crear nueva lección',                            'Profesor'],
        ['DELETE', '/api/lessons/:id',         'Eliminar lección del catálogo',                  'Profesor'],
        ['GET',    '/api/progress',            'Progreso del alumno autenticado',                'Autenticado'],
        ['GET',    '/api/progress/stats',      'Estadísticas: total lecciones completadas',      'Autenticado'],
        ['GET',    '/api/progress/all',        'Historial de todos los alumnos',                 'Profesor'],
        ['POST',   '/api/progress/complete',   'Completar lección y guardar score del quiz',     'Autenticado'],
        ['DELETE', '/api/progress/:lessonId',  'Desmarcar lección como completada',              'Autenticado'],
        ['GET',    '/api/ratings',             'Valoraciones del usuario autenticado',           'Autenticado'],
        ['POST',   '/api/ratings',             'Publicar valoración (estrellas + comentario)',   'Autenticado'],
        ['GET',    '/api/ratings/stats',       'Media global de valoraciones',                   'Autenticado'],
    ],
    headers=['Método', 'Ruta', 'Descripción', 'Acceso']
)

divider(doc)

# ── 6. SEGURIDAD ─────────────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '6. Seguridad y autenticación', '1A3A5C')

body(doc, '6.1 Sistema JWT', color=PRIMARY, size=12, bold=True)
body(doc,
    'La autenticación utiliza JSON Web Tokens firmados con un secreto definido en variables de entorno. '
    'El token incluye el ID del usuario, su email, su rol y la fecha de expiración (7 días).', size=10.5)
doc.add_paragraph()
code_block(doc,
'''// Payload del JWT
{
  "sub":   "6a25f6b770670b52163a1b25",  // userId
  "email": "alumno@codex.com",
  "rol":   "alumno",
  "iat":   1718000000,                  // issued at
  "exp":   1718604800                   // expira en 7 días
}''')

doc.add_paragraph()
body(doc, '6.2 Flujo de autenticación', color=PRIMARY, size=12, bold=True)
bullet(doc, 'El usuario introduce email y contraseña en el formulario de login.')
bullet(doc, 'El backend compara la contraseña con el hash bcrypt almacenado.')
bullet(doc, 'Si es correcto, se genera un JWT firmado y se devuelve al cliente.')
bullet(doc, 'El frontend guarda el token en localStorage.')
bullet(doc, 'El interceptor Angular adjunta el token en cada petición HTTP posterior.')
bullet(doc, 'El JwtAuthGuard de NestJS verifica el token en cada endpoint protegido.')
bullet(doc, 'Si el token expira o es inválido, el servidor devuelve 401 y el interceptor redirige al login.')

doc.add_paragraph()
body(doc, '6.3 Control de roles', color=PRIMARY, size=12, bold=True)
body(doc,
    'El rol está incluido en el JWT payload. Tanto el frontend (guards de ruta Angular) como el backend '
    '(comprobación en cada endpoint sensible) verifican el rol. Un alumno no puede acceder a ninguna '
    'ruta de /profesor/* ni ejecutar operaciones de gestión en la API.', size=10.5)

doc.add_paragraph()
body(doc, '6.4 Otras medidas de seguridad', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Contraseñas nunca almacenadas en texto plano — siempre hasheadas con bcrypt (salt 10).')
bullet(doc, 'El campo password se excluye de todas las respuestas de la API con .select("-password").')
bullet(doc, 'Variables de entorno (.env) para secretos — MONGO_URL y JWT_SECRET fuera del código fuente.')
bullet(doc, 'Validación global de DTOs con ValidationPipe (whitelist: true) — se rechazan campos no esperados.')
bullet(doc, 'ETags deshabilitados en Express para evitar respuestas 304 que dejaban el cliente en estado inconsistente.')
bullet(doc, 'Flag redirectingToLogin en el interceptor para evitar race conditions con peticiones paralelas.')

divider(doc)

# ── 7. FUNCIONALIDADES ───────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '7. Funcionalidades de la plataforma', '1A3A5C')

body(doc, '7.1 Registro e inicio de sesión', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Formulario de registro con nombre, email y contraseña. El rol por defecto es "alumno".')
bullet(doc, 'Formulario de login con validación y mensaje de error en credenciales incorrectas.')
bullet(doc, 'Redirección automática al dashboard correspondiente según el rol del usuario.')
bullet(doc, 'Guard de autenticación que protege todas las rutas internas y redirige al login si no hay sesión.')

doc.add_paragraph()
body(doc, '7.2 Catálogo y lecciones (alumno)', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Catálogo con todas las lecciones disponibles, filtrable por categoría (HTML, CSS, JavaScript) y nivel.')
bullet(doc, 'Tarjeta de lección con título, descripción, nivel, categoría e indicador de completado.')
bullet(doc, 'Vista de detalle con secciones de teoría y bloques de código resaltados con Prism.js.')
bullet(doc, 'Quiz al final de cada lección con preguntas de opción múltiple y resultado inmediato.')
bullet(doc, 'Al completar el quiz, el progreso y la puntuación se guardan en el servidor.')
bullet(doc, 'Sistema de valoración (1-5 estrellas + comentario) que se muestra en el panel del profesor.')

doc.add_paragraph()
body(doc, '7.3 Dashboard y progreso del alumno', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Panel principal con barra de progreso global (lecciones completadas / total).')
bullet(doc, 'Lista de últimas lecciones completadas con fecha.')
bullet(doc, 'Sistema de logros (achievements) desbloqueados automáticamente por hitos.')
bullet(doc, 'Perfil de usuario editable (nombre).')
bullet(doc, 'Página de progreso detallado con historial completo de lecciones completadas.')

doc.add_paragraph()
body(doc, '7.4 Panel del profesor', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Dashboard con tarjetas resumen: alumnos activos, lecciones creadas, tests completados, valoración media.')
bullet(doc, 'Lista de alumnos con nombre, email, lecciones completadas y porcentaje de progreso.')
bullet(doc, 'Formulario para añadir nuevos alumnos (nombre, email, contraseña) directamente desde la plataforma.')
bullet(doc, 'Gestión del catálogo de lecciones: vista completa con opción de eliminar.')
bullet(doc, 'Historial de tests completados: quién completó qué lección, puntuación y fecha.')
bullet(doc, 'Lista de valoraciones recibidas con estrellas, comentario y media global.')

doc.add_paragraph()
body(doc, '7.5 Gestión de lecciones (formulario de creación)', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Formulario estructurado para crear lecciones con título, descripción, nivel y categoría.')
bullet(doc, 'Secciones de contenido con campo de texto y bloque de código opcional.')
bullet(doc, 'Preguntas de quiz con hasta 4 opciones y selección de la respuesta correcta.')
bullet(doc, 'La lección se almacena en MongoDB y aparece inmediatamente en el catálogo.')

divider(doc)

# ── 8. RETOS TÉCNICOS ───────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '8. Retos técnicos y soluciones', '1A3A5C')

challenges = [
    (
        'CORS y proxy de desarrollo',
        'Las peticiones directas del navegador al backend (localhost:3000) quedaban colgadas indefinidamente '
        'sin completar ni dar error, lo que impedía cargar cualquier dato.',
        'Se configuró un proxy en Angular (proxy.conf.json) que intercepta todas las rutas /api/* y las '
        'reenvía al backend a través del propio servidor de desarrollo. Se añadió también el prefijo global '
        '/api en NestJS para que las URLs coincidan sin reescritura de rutas.'
    ),
    (
        'Change detection en Angular 21 con Vite',
        'Las respuestas HTTP completaban correctamente (200 OK visto en Network), los datos llegaban '
        'al componente, pero la vista permanecía congelada mostrando el estado de carga.',
        'Angular 21 utiliza Vite como servidor de desarrollo. Zone.js no parchea el Fetch API de Vite, '
        'por lo que las respuestas asíncronas se resuelven fuera de la zona de Angular y no disparan el '
        'ciclo de detección de cambios. Solución: inyectar ChangeDetectorRef y llamar a detectChanges() '
        'manualmente tras cada actualización de estado.'
    ),
    (
        'Caché HTTP 304 Not Modified',
        'Express genera ETags automáticamente para todas las respuestas. El navegador enviaba la cabecera '
        'If-None-Match con el ETag conocido. El servidor respondía 304 sin cuerpo. Angular recibía un '
        '304 que no disparaba el callback next() del observable, dejando loading = true para siempre.',
        'Se deshabilitó la generación de ETags en Express con app.getHttpAdapter().getInstance().set("etag", false). '
        'Esto fuerza al servidor a devolver siempre 200 con el cuerpo completo.'
    ),
    (
        'DNS intermitente con MongoDB Atlas en macOS',
        'El backend arrancaba pero las consultas a MongoDB fallaban con ENOTFOUND para los hostnames '
        'de los shards del cluster. El error se producía de forma intermitente.',
        'Node.js intentaba resolver los hostnames de MongoDB vía IPv6, que fallaba en el DNS de la red local. '
        'Solución: añadir family: 4 en las opciones de conexión de Mongoose para forzar IPv4. También se '
        'añadieron timeouts de conexión (serverSelectionTimeoutMS: 8000) para que los fallos sean '
        'rápidos y no bloqueen indefinidamente.'
    ),
    (
        'Race condition en el interceptor JWT',
        'Al hacer 3 peticiones simultáneas que fallaban con 401 (token expirado), el interceptor '
        'ejecutaba localStorage.clear() y navigate("/login") tres veces en paralelo, causando '
        'estado impredecible en la aplicación.',
        'Se añadió una variable booleana redirectingToLogin a nivel de módulo (fuera de la función '
        'del interceptor). La primera petición con 401 establece el flag a true, ejecuta la redirección '
        'y lo resetea a false tras navegar. Las peticiones 2 y 3 ven el flag a true y no disparan '
        'ninguna acción adicional.'
    ),
    (
        'Invalidación de tokens tras re-seed de la base de datos',
        'Al re-ejecutar el script de seed, todos los usuarios son eliminados y recreados con nuevos '
        'IDs de MongoDB. Los JWT existentes en el navegador contenían el ID antiguo, por lo que el '
        'servidor no encontraba el usuario y lanzaba 401.',
        'Se añadió un fallback en la JWT Strategy de Passport: si el usuario no se encuentra por ID (payload.sub), '
        'se intenta una segunda búsqueda por email (payload.email). Como el email no cambia entre re-seeds, '
        'el usuario se identifica correctamente y la sesión sigue siendo válida.'
    ),
]

for i, (title, problem, solution) in enumerate(challenges, 1):
    body(doc, f'8.{i} {title}', color=ACCENT, size=11, bold=True)

    body(doc, 'Problema:', color=MUTED, size=9, bold=True)
    body(doc, problem, size=10)

    body(doc, 'Solución:', color=GREEN, size=9, bold=True)
    body(doc, solution, size=10)
    doc.add_paragraph()

divider(doc)

# ── 9. ESTRUCTURA DE CARPETAS ────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '9. Guía de despliegue y uso', '1A3A5C')

body(doc, '9.1 Requisitos previos', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Node.js v18 o superior (probado con v24)')
bullet(doc, 'npm v9 o superior')
bullet(doc, 'Cuenta en MongoDB Atlas con un cluster M0 (gratuito)')
bullet(doc, 'Git para clonar el repositorio')

doc.add_paragraph()
body(doc, '9.2 Instalación y configuración', color=PRIMARY, size=12, bold=True)
code_block(doc,
'''# 1. Clonar el repositorio
git clone https://github.com/JRGamer234/Codex-FCT26.git
cd Codex-FCT26

# 2. Instalar dependencias (frontend y backend)
npm install

# 3. Configurar variables de entorno del backend
# Editar backend/.env:
MONGO_URL=mongodb+srv://<usuario>:<password>@<cluster>.mongodb.net/codex
JWT_SECRET=tu-secreto-seguro
PORT=3000

# 4. Poblar la base de datos con datos de prueba
cd backend && npx ts-node src/seed.ts && cd ..

# 5. Arrancar frontend y backend simultáneamente
npm run dev''')

doc.add_paragraph()
body(doc, '9.3 Credenciales de prueba (tras ejecutar seed)', color=PRIMARY, size=12, bold=True)
info_table(doc,
    data=[
        ['Profesor Demo', 'profesor@codex.com', '123456', 'profesor'],
        ['Alumno Demo',   'alumno@codex.com',   '123456', 'alumno'],
        ['Alex',          'alex@codex.com',     '123456', 'alumno'],
        ['Jorge',         'jorge@codex.com',    '123456', 'alumno'],
        ['Mario',         'mario@codex.com',    '123456', 'alumno'],
        ['Itziar',        'itziar@codex.com',   '123456', 'alumno'],
    ],
    headers=['Nombre', 'Email', 'Contraseña', 'Rol']
)

doc.add_paragraph()
body(doc, '9.4 Scripts disponibles', color=PRIMARY, size=12, bold=True)
info_table(doc,
    data=[
        ['npm run dev',              'Raíz',     'Arranca frontend y backend simultáneamente'],
        ['npm run frontend',         'Raíz',     'Solo frontend en localhost:4200'],
        ['npm run start:dev',        'backend/', 'Backend con hot-reload (watch mode)'],
        ['npx ts-node src/seed.ts',  'backend/', 'Poblar/reiniciar la base de datos'],
        ['ng build',                 'frontend/', 'Build de producción'],
    ],
    headers=['Comando', 'Directorio', 'Descripción']
)

divider(doc)

# ── 10. CONCLUSIONES ─────────────────────────────────────────────────────────
doc.add_page_break()
section_title_box(doc, '10. Conclusiones', '0E4F7A')

body(doc,
    'Codex es una plataforma educativa full-stack completamente funcional que demuestra la integración '
    'de tecnologías modernas para el desarrollo web. El proyecto abarca desde la autenticación segura '
    'hasta la gestión de contenido educativo, pasando por la experiencia de usuario tanto para alumnos '
    'como para profesores.',
    size=10.5)

doc.add_paragraph()
body(doc, 'Logros principales', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Plataforma educativa funcional con 7 lecciones reales de HTML, CSS y JavaScript.')
bullet(doc, 'Arquitectura full-stack con separación clara de responsabilidades (Angular + NestJS + MongoDB).')
bullet(doc, 'Autenticación JWT segura con control de acceso por rol en frontend y backend.')
bullet(doc, 'Sistema de seguimiento de progreso con puntuación de quizzes y historial.')
bullet(doc, 'Panel de profesor completo: gestión de alumnos, lecciones, tests y valoraciones.')
bullet(doc, 'Resolución de 6 problemas técnicos complejos durante el desarrollo.')
bullet(doc, 'Código versionado en GitHub con historial de commits estructurado.')

doc.add_paragraph()
body(doc, 'Trabajo futuro', color=PRIMARY, size=12, bold=True)
bullet(doc, 'Editor visual de lecciones (WYSIWYG) para el profesor.')
bullet(doc, 'Analíticas con gráficos de progreso histórico por alumno.')
bullet(doc, 'Sistema de mensajería entre alumno y profesor.')
bullet(doc, 'Despliegue en producción (Vercel / Railway / Render).')
bullet(doc, 'Tests unitarios e integración para garantizar la calidad del código.')
bullet(doc, 'Soporte para contenido multimedia (vídeos, ejercicios interactivos de código).')

doc.add_paragraph()
body(doc, 'github.com/JRGamer234/Codex-FCT26', color=PRIMARY, size=11, bold=True)

# ── GUARDAR ──────────────────────────────────────────────────────────────────
output_path = '/Users/itzzc/Documents/tfg/Codex-FCT26/docs/Codex_FCT26_Documentacion.docx'
doc.save(output_path)
print(f'Documento guardado: {output_path}')
